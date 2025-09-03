import requests
from bs4 import BeautifulSoup
from docx import Document

# URL of the article
url = "https://www.elmundo.es/cronica/2024/07/29/66a3f99be4d4d8506e8b456d.html"

# Function to fetch and parse the article content
def fetch_article_content(url):
    headers = {
        "User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/91.0.4472.124 Safari/537.36"
    }
    response = requests.get(url, headers=headers)
    soup = BeautifulSoup(response.content, 'html.parser')
    
    # Update the class name according to the structure of the webpage
    article_body = soup.find('div', class_='ue-c-article__body ue-l-article__body')
    if article_body:
        paragraphs = article_body.find_all('p')
        article_text = "\n\n".join([para.get_text() for para in paragraphs])
        return article_text
    return None

# Fetch the article content
article_content = fetch_article_content(url)

# Create a Word document and add the article content
if article_content:
    doc = Document()
    doc.add_heading('Las "identidades" del comprador de una autocaravana condenado a pagar 20.000 euros por atacar "el honor" del vendedor', level=1)
    doc.add_paragraph(article_content)
    
    # Save the document
    file_path = "Articulo_ElMundo.docx"
    doc.save(file_path)
    print(f"Documento guardado en: {file_path}")
else:
    print("No se pudo extraer el contenido del artículo.")

